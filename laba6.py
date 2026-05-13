import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from geometry import Rectangle, Triangle, Trapezoid
import openpyxl
from docx import Document
from datetime import datetime

# ========== GUI на Tkinter  ==========
class GeometryApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Геометрические фигуры — Вариант 8")
        self.root.geometry("500x550")
        self.root.configure(bg='#f0f0f0')

        self.figure_var = tk.StringVar(value="Прямоугольник")

        # Заголовок
        title = tk.Label(root, text="Вариант 8: Прямоугольник / Треугольник / Трапеция",
                         font=("Arial", 12, "bold"), bg='#f0f0f0')
        title.pack(pady=10)

        # Рамка с выбором фигуры
        frame_fig = tk.LabelFrame(root, text="Выберите фигуру", bg='#f0f0f0', font=("Arial", 10, "bold"))
        frame_fig.pack(pady=5, padx=20, fill="x")

        tk.Radiobutton(frame_fig, text="📐 Прямоугольник", variable=self.figure_var,
                       value="Прямоугольник", command=self.update_inputs, bg='#f0f0f0').pack(anchor="w", padx=20, pady=2)
        tk.Radiobutton(frame_fig, text="📐 Треугольник", variable=self.figure_var,
                       value="Треугольник", command=self.update_inputs, bg='#f0f0f0').pack(anchor="w", padx=20, pady=2)
        tk.Radiobutton(frame_fig, text="📐 Трапеция", variable=self.figure_var,
                       value="Трапеция", command=self.update_inputs, bg='#f0f0f0').pack(anchor="w", padx=20, pady=2)

        # Рамка для ввода параметров
        self.frame_inputs = tk.LabelFrame(root, text="Параметры фигуры", bg='#f0f0f0', font=("Arial", 10, "bold"))
        self.frame_inputs.pack(pady=10, padx=20, fill="x")

        self.entries = {}
        self.update_inputs()

        # Кнопка расчёта
        self.calc_btn = tk.Button(root, text="▶ РАССЧИТАТЬ", command=self.calculate,
                                  bg="#4CAF50", fg="white", font=("Arial", 11, "bold"), padx=20, pady=5)
        self.calc_btn.pack(pady=10)

        # Рамка для результата
        frame_result = tk.LabelFrame(root, text="Результат", bg='#f0f0f0', font=("Arial", 10, "bold"))
        frame_result.pack(pady=10, padx=20, fill="both", expand=True)

        self.result_text = tk.Text(frame_result, height=8, width=50, bg="#ffffe0", font=("Arial", 10))
        self.result_text.pack(padx=10, pady=10, fill="both", expand=True)

        # Кнопки сохранения
        btn_frame = tk.Frame(root, bg='#f0f0f0')
        btn_frame.pack(pady=10)

        tk.Button(btn_frame, text="📎 Сохранить в Excel", command=self.save_excel,
                  bg="#2196F3", fg="white", padx=10).pack(side="left", padx=10)
        tk.Button(btn_frame, text="📄 Сохранить в Word", command=self.save_word,
                  bg="#2196F3", fg="white", padx=10).pack(side="left", padx=10)

    def update_inputs(self):
        # Очищаем старые поля
        for widget in self.frame_inputs.winfo_children():
            widget.destroy()
        self.entries.clear()

        figure = self.figure_var.get()

        if figure == "Прямоугольник":
            labels = ["a (сторона)", "b (сторона)"]
        elif figure == "Треугольник":
            labels = ["a (сторона)", "b (сторона)", "c (сторона)"]
        else:  # Трапеция
            labels = ["a (основание)", "b (основание)", "h (высота)"]

        for i, label in enumerate(labels):
            tk.Label(self.frame_inputs, text=label, bg='#f0f0f0', font=("Arial", 10)).grid(row=i, column=0, padx=10, pady=5, sticky="w")
            entry = tk.Entry(self.frame_inputs, width=15, font=("Arial", 10))
            entry.grid(row=i, column=1, padx=10, pady=5)
            self.entries[label] = entry

    def get_figure(self):
        figure = self.figure_var.get()
        try:
            if figure == "Прямоугольник":
                a = float(self.entries["a (сторона)"].get())
                b = float(self.entries["b (сторона)"].get())
                return Rectangle(a, b), figure, {"a": a, "b": b}
            elif figure == "Треугольник":
                a = float(self.entries["a (сторона)"].get())
                b = float(self.entries["b (сторона)"].get())
                c = float(self.entries["c (сторона)"].get())
                if a + b <= c or a + c <= b or b + c <= a:
                    raise ValueError("Такой треугольник не существует")
                return Triangle(a, b, c), figure, {"a": a, "b": b, "c": c}
            elif figure == "Трапеция":
                a = float(self.entries["a (основание)"].get())
                b = float(self.entries["b (основание)"].get())
                h = float(self.entries["h (высота)"].get())
                return Trapezoid(a, b, h), figure, {"a": a, "b": b, "h": h}
        except ValueError as e:
            messagebox.showerror("Ошибка", str(e))
            return None, None, None

    def calculate(self):
        fig, figure_name, params = self.get_figure()
        if fig is None:
            return

        area = fig.area()
        perimeter = fig.perimeter()

        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, f"{fig.info()}\n\n")
        self.result_text.insert(tk.END, f"📐 Площадь: {area:.2f}\n")
        self.result_text.insert(tk.END, f"📏 Периметр: {perimeter:.2f}\n")
        self.result_text.insert(tk.END, f"\n✅ Результат готов")

    def save_excel(self):
        fig, figure_name, params = self.get_figure()
        if fig is None:
            return
        area = fig.area()
        perimeter = fig.perimeter()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Результаты"
        ws.append(["Фигура", "Параметры", "Площадь", "Периметр", "Дата"])
        ws.append([figure_name, str(params), area, perimeter, str(datetime.now())])

        filepath = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                filetypes=[("Excel files", "*.xlsx")])
        if filepath:
            wb.save(filepath)
            messagebox.showinfo("Сохранено", f"Файл сохранён:\n{filepath}")

    def save_word(self):
        fig, figure_name, params = self.get_figure()
        if fig is None:
            return
        area = fig.area()
        perimeter = fig.perimeter()

        doc = Document()
        doc.add_heading("Отчёт по геометрической фигуре", level=1)
        doc.add_paragraph(f"Фигура: {figure_name}")
        doc.add_paragraph(f"Параметры: {params}")
        doc.add_paragraph(f"Площадь: {area:.2f}")
        doc.add_paragraph(f"Периметр: {perimeter:.2f}")
        doc.add_paragraph(f"Дата расчёта: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        filepath = filedialog.asksaveasfilename(defaultextension=".docx",
                                                filetypes=[("Word files", "*.docx")])
        if filepath:
            doc.save(filepath)
            messagebox.showinfo("Сохранено", f"Файл сохранён:\n{filepath}")

if __name__ == "__main__":
    root = tk.Tk()
    app = GeometryApp(root)
    root.mainloop()
